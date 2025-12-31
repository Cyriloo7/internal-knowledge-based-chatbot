# client/src/components/indexer.py

import io
import os
from PIL import Image
from langchain_core.documents import Document
from dotenv import load_dotenv
import fitz  # PyMuPDF
from langchain_community.document_loaders import PyPDFLoader, TextLoader, UnstructuredPowerPointLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_chroma import Chroma
import google.generativeai as genai


class Indexer:
    def __init__(self, file_path):
        load_dotenv()
        genai.configure(api_key=os.environ["GOOGLE_API_KEY"])
        self.file_path = file_path
        self.file_ext = os.path.splitext(file_path)[1].lower()
        self.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
        self.vision_model = genai.GenerativeModel('models/gemini-2.5-flash')

    def extract_images_from_doc(self, pdf_path):
        """Extract images from PDF files"""
        if self.file_ext != '.pdf':
            return []
        
        doc = fitz.open(pdf_path)
        images = []
        for page_index in range(len(doc)):
            page = doc[page_index]
            for img in page.get_images(full=True):
                xref = img[0]
                base = doc.extract_image(xref)
                image_bytes = base["image"]
                images.append((page_index+1, image_bytes))
        return images
    
    def image_to_document(self, image_bytes, page_num):
        """Convert image to document with OCR and caption"""
        image = Image.open(io.BytesIO(image_bytes))

        # OCR
        import pytesseract
        ocr_text = pytesseract.image_to_string(image).strip()

        # Gemini caption
        caption_prompt = """
        Generate a concise technical caption for this image.
        Describe diagrams, architecture, flow, components, and relationships.
        """
        caption = self.vision_model.generate_content([caption_prompt, image])
        caption = caption.text.strip()

        combined_text = f"""
        IMAGE CAPTION:
        {caption}
        OCR TEXT:
        {ocr_text}
        """.strip()

        return Document(
            page_content=combined_text,
            metadata={
                "type": "image",
                "source": self.file_path,
                "page_num": page_num,
            }
        )
    
    def load_document(self):
        """Load document based on file extension"""
        if self.file_ext == '.pdf':
            loader = PyPDFLoader(self.file_path)
            pages = loader.load()
            return pages
        elif self.file_ext == '.txt':
            loader = TextLoader(self.file_path, encoding='utf-8')
            pages = loader.load()
            return pages
        elif self.file_ext in ['.ppt', '.pptx']:
            loader = UnstructuredPowerPointLoader(self.file_path)
            pages = loader.load()
            return pages
        elif self.file_ext in ['.doc', '.docx']:
            # Use python-docx for DOCX files
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(self.file_path)
                full_text = []
                for para in doc.paragraphs:
                    full_text.append(para.text)
                
                # Combine all paragraphs into a single document
                full_text_str = '\n'.join(full_text)
                return [Document(page_content=full_text_str, metadata={"source": self.file_path})]
            except ImportError:
                raise ImportError("python-docx is required for DOCX files. Install it with: pip install python-docx")
        else:
            raise ValueError(f"Unsupported file format: {self.file_ext}")
    
    def index_document(self, level):
        """Index a document into the vector store"""
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"The file {self.file_path} does not exist.")

        # Load document based on file type
        try:
            pages = self.load_document()
        except Exception as e:
            raise RuntimeError(f"Failed to load file: {e}")
        
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=3000, chunk_overlap=500)
        docs_split = text_splitter.split_documents(pages)
        for doc in docs_split:
            doc.metadata.update({"type": "text", "source": self.file_path})
    
        # Extract images only from PDF files
        image_doc = []
        if self.file_ext == '.pdf':
            for pagenum, image_bytes in self.extract_images_from_doc(self.file_path):
                try:
                    image_doc.append(self.image_to_document(image_bytes, pagenum))
                except Exception as e:
                    print(f"Error processing image on page {pagenum}: {e}")

        vectorstorepath = "chroma_vectorstore"
        collection_name = level

        path = os.path.join(os.getcwd(), vectorstorepath, collection_name)
        os.makedirs(path, exist_ok=True)

        import hashlib

        def doc_id(doc):
            content = doc.page_content
            source = doc.metadata.get("source", "")
            h = hashlib.sha256(f"{source}:{content}".encode()).hexdigest()
            return h

        all_docs = docs_split + image_doc
        ids = [doc_id(doc) for doc in all_docs]

        vector_store = Chroma(
            collection_name=collection_name,
            embedding_function=self.embeddings,
            persist_directory=path,
        )
        existing = vector_store._collection.get(include=[])
        existing_ids = set(existing["ids"])

        new_docs = []
        new_ids = []

        for doc, id_ in zip(all_docs, ids):
            if id_ not in existing_ids:
                new_docs.append(doc)
                new_ids.append(id_)

        try:
            if new_docs:
                vector_store.add_documents(new_docs, ids=new_ids)
                print(f"✅ Indexed {len(new_docs)} new documents into collection '{collection_name}'")
            else:
                print("✅ No new documents to add")

        except Exception as e:
            raise RuntimeError(f"Failed to create or persist Chroma vector store: {e}")
